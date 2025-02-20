import flet as ft
import os
from time import sleep
from docx import Document


def main(page: ft.Page):
    page.title = "Cadastro de Pacientes"
    page.vertical_alignment = ft.MainAxisAlignment.CENTER
    page.horizontal_alignment = ft.MainAxisAlignment.CENTER
    page.padding = 10
    page.theme_mode = 'light'
    page.spacing = 15
    page.update()

    # Banner para aviso os erros

    def close_banner(e):
        page.close(banner)
        # page.add(ft.Text("Action clicked: " + e.control.text))

    action_button_style = ft.ButtonStyle(color=ft.Colors.BLUE)
    banner = ft.Banner(
        bgcolor=ft.Colors.AMBER_100,
        leading=ft.Icon(ft.Icons.WARNING, color=ft.Colors.RED, size=40),
        content=ft.Text(
            value="Oops, você precisa digitar todos os campos para efetuar o cadastro ",
            color=ft.Colors.BLACK,
            size=18,
        ),
        actions=[
            ft.TextButton(text="Fechar", style=action_button_style,
                          on_click=close_banner),
        ],
    )

    def close_banner1(e):
        page.close(banner1)
    banner1 = ft.Banner(
        bgcolor=ft.Colors.AMBER_100,
        leading=ft.Icon(ft.Icons.WARNING, color=ft.Colors.RED, size=40),
        content=ft.Text(
            value="Oops, Digite Nome, e-mail ou CPF para consulta ",
            color=ft.Colors.BLACK,
            size=18,
        ),
        actions=[
            ft.TextButton(text="Fechar", style=action_button_style,
                          on_click=close_banner1),
        ],
    )
    page.update()

    # Campos de entrada
    nome = ft.TextField(label="Nome", width=500)
    data_nascimento = ft.TextField(label="Data de Nascimento", width=300)
    email = ft.TextField(label="E-mail", width=500, icon='EMAIL')
    telefone = ft.TextField(label="Telefone", width=300, icon='PHONE')
    cpf = ft.TextField(label="CPF", width=300)
    endereco = ft.TextField(label="Endereço", width=500)
    informacao_importante = ft.TextField(
        label="Informação Importante", width=810,)

    # Ajuste a altura conforme necessário
    mensagem = ft.Text("", color=ft.colors.RED_400,
                       size=22, text_align='CENTER')

    # Layout com os campos organizados
    linha_nome_data = ft.Row([nome, data_nascimento],
                             alignment=ft.MainAxisAlignment.CENTER)
    linha_email_telefone = ft.Row(
        [email, telefone], alignment=ft.MainAxisAlignment.CENTER)
    linha_cpf_endereco = ft.Row(
        [cpf, endereco], alignment=ft.MainAxisAlignment.CENTER)

    # Função para salvar em um arquivo .docx
    def salvar_em_doc(paciente):
        filename = f'{paciente["nome"]}.doc'
        doc = Document()
        doc.add_heading(paciente["nome"], level=1)
        doc.add_paragraph(f"Nome: {paciente['nome']}")
        doc.add_paragraph(f"Data de Nascimento: {paciente['data_nascimento']}")
        doc.add_paragraph(f"E-mail: {paciente['email']}")
        doc.add_paragraph(f"Telefone: {paciente['telefone']}")
        doc.add_paragraph(f"CPF: {paciente['cpf']}")
        doc.add_paragraph(f"Endereço: {paciente['endereco']}")
        doc.add_paragraph(
            f"Informação Importante: {paciente['informacao_importante']}")
        doc.add_paragraph("-" * 50)
        doc.save(filename)

    # Função de cadastro
    def cadastrar(e):
        if not all([nome.value, data_nascimento.value, email.value, telefone.value, cpf.value, endereco.value, informacao_importante.value]):
            page.open(banner)
            # mensagem.value = "Por Favor, Preencha todos os campos!"
            # mensagem.update()
        else:
            paciente = {
                "nome": nome.value,
                "data_nascimento": data_nascimento.value,
                "email": email.value,
                "telefone": telefone.value,
                "cpf": cpf.value,
                "endereco": endereco.value,
                "informacao_importante": informacao_importante.value
            }
            mensagem.value = "Paciente cadastrado com sucesso!"
            mensagem.update()
        salvar_em_doc(paciente)

    page.update()

    # Função para limpar os campos
    def limpar(e):
        nome.value = ""
        data_nascimento.value = ""
        email.value = ""
        telefone.value = ""
        cpf.value = ""
        endereco.value = ""
        informacao_importante.value = ""
        mensagem.value = ""
        page.update()

    # Função para consultar paciente
    def consultar(e):
        termo = nome.value.strip() or cpf.value.strip(
        ) or email.value.strip() or telefone.value.strip()
        if not termo:
            page.open(banner1)
            # mensagem.value = "Digite Nome, CPF, E-mail ou Telefone para consulta!"
        else:
            filename = f"{nome.value}.doc"
            if os.path.exists(filename):
                doc = Document(filename)
                dados = {p.text.split(": ")[0]: p.text.split(
                    ": ")[1] for p in doc.paragraphs if ": " in p.text}
                nome.value = dados.get("Nome", "")
                data_nascimento.value = dados.get("Data de Nascimento", "")
                email.value = dados.get("E-mail", "")
                telefone.value = dados.get("Telefone", "")
                cpf.value = dados.get("CPF", "")
                endereco.value = dados.get("Endereço", "")
                informacao_importante.value = dados.get(
                    "Informação Importante", "")
                mensagem.value = "Paciente encontrado!"
            else:
                mensagem.value = "Paciente não encontrado."
        page.update()

    # Botões
    btn_cadastrar = ft.ElevatedButton("Cadastrar", on_click=cadastrar)
    btn_limpar = ft.ElevatedButton("Limpar", on_click=limpar)
    btn_consultar = ft.ElevatedButton("Consultar", on_click=consultar)

    page.add(
        linha_nome_data,
        linha_email_telefone,
        linha_cpf_endereco,
        informacao_importante,
        mensagem,
        banner,
        ft.Row([btn_cadastrar, btn_limpar, btn_consultar],
               alignment=ft.MainAxisAlignment.CENTER)
    )


ft.app(target=main)
